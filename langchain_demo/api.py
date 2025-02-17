from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from typing import List
import os
import shutil
from datetime import datetime
import logging
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_openai import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from docx import Document as DocxDocument
import PyPDF2
import pdfplumber
import pandas as pd
from langchain.schema import Document
from dotenv import load_dotenv

# Tải các biến môi trường từ file .env
load_dotenv()

# Khởi tạo FastAPI app
app = FastAPI(title="Document Upload API")

# Cấu hình CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Cấu hình
UPLOAD_DIR = "data/documents"
CHROMA_PERSIST_DIR = "data/chroma_db"
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# Cấu hình logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def ensure_directories():
    """Đảm bảo các thư mục cần thiết tồn tại"""
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    os.makedirs(CHROMA_PERSIST_DIR, exist_ok=True)

def process_file(file_path: str) -> str:
    """Xử lý file và trích xuất nội dung"""
    file_extension = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
                
        elif file_extension in ['.doc', '.docx']:
            doc = DocxDocument(file_path)
            return ' '.join([paragraph.text for paragraph in doc.paragraphs])
            
        elif file_extension == '.pdf':
            with pdfplumber.open(file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    text += page.extract_text()
                return text
            
        else:
            raise ValueError(f"Không hỗ trợ định dạng file {file_extension}")
            
    except Exception as e:
        logger.error(f"Lỗi khi xử lý file {file_path}: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing file: {e}")

from langchain.schema import Document  # Đảm bảo đã import Document từ langchain.schema

def update_chroma_db(texts: List[str], sources: List[dict]):
    """Cập nhật ChromaDB với documents mới và metadata (không bao gồm source)"""
    try:
        embeddings = OpenAIEmbeddings(
            api_key=OPENAI_API_KEY,
            model="text-embedding-3-small"
        )
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2000,
            chunk_overlap=200
        )
        
        # Tạo documents với metadata, không bao gồm source
        documents = []
        for text, source in zip(texts, sources):
            splits = text_splitter.split_text(text)
            for split in splits:
                document = Document(
                    page_content=split,  # Nội dung của chunk
                    metadata={           # Metadata của document
                        "university_id": source["university_id"],  # Chỉ giữ lại university_id
                        "timestamp": datetime.now().isoformat()  # Thêm thời gian tải lên
                    }
                )
                documents.append(document)
        
        # Lưu vào ChromaDB
        vectorstore = Chroma.from_documents(
            documents=documents,
            embedding=embeddings,
            persist_directory=CHROMA_PERSIST_DIR
        )
        vectorstore.persist()
        
        return len(documents)
    except Exception as e:
        logger.error(f"Lỗi khi cập nhật ChromaDB: {e}")
        raise

    except Exception as e:
        logger.error(f"Lỗi khi cập nhật ChromaDB: {e}")
        raise


@app.post("/upload/")
async def upload_files(
    files: List[UploadFile] = File(...),
    university_id: str = Form(...)  # Chỉ nhận university_id
):
    """API endpoint để upload files kèm theo ID trường"""
    ensure_directories()
    
    try:
        processed_texts = []
        file_sources = []
        
        for file in files:
            # Tạo tên file unique
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            file_name = f"{timestamp}_{file.filename}"
            file_path = os.path.join(UPLOAD_DIR, file_name)
            
            # Lưu file
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            
            # Xử lý file
            text = process_file(file_path)
            processed_texts.append(text)
            file_sources.append({"university_id": university_id})  # Chỉ lưu lại university_id vào metadata
            
            logger.info(f"Đã xử lý file: {file_name} từ trường {university_id}")
        
        # Cập nhật ChromaDB
        num_chunks = update_chroma_db(processed_texts, file_sources)
        
        return {
            "message": "Upload thành công",
            "files_processed": len(files),
            "chunks_created": num_chunks,
            "university_id": university_id
        }
        
    except Exception as e:
        logger.error(f"Lỗi trong quá trình upload: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/health")
async def health_check():
    """API endpoint để kiểm tra health"""
    return {"status": "healthy"}
