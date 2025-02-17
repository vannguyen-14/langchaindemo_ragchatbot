from langchain.chains import RetrievalQA
from langchain_openai import OpenAI, OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from langchain.utilities import SerpAPIWrapper
from keybert import KeyBERT
from pinecone import Pinecone as PineconeClient
import logging
from typing import List, Optional
import os
import json
from datetime import datetime
import chromadb
from langchain_core.documents import Document
from docx import Document as DocxDocument
import PyPDF2
import pandas as pd
from io import BytesIO

# Định nghĩa đường dẫn
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
LOCAL_DATA_DIR = os.path.join(BASE_DIR, "data", "documents")

# Thiết lập logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# Cấu hình
OPENAI_API_KEY = "your_openai_api_key"
SERP_API_KEY = "062caddf2fd2f14c079e2a99e38afc4b3733ba36b9afa1d755d8cf202db6f469"
SESSION_DIR = "data/sessions"  # Thư mục lưu session

# Thêm config cho Chroma
CHROMA_PERSIST_DIR = "data/chroma_db"  # Thư mục lưu vector store

# System prompt
SYSTEM_PROMPT = """Bạn là một trợ lý tư vấn tuyển sinh thông minh. Khi trả lời:
1. Hãy tóm tắt thông tin một cách ngắn gọn, dễ hiểu
2. Trình bày thông tin theo dạng có cấu trúc với các bullet points
3. Nếu thông tin từ web, hãy trích dẫn nguồn và đưa vào dạng link
4. Tập trung vào những thông tin quan trọng và liên quan nhất
5. Sử dụng ngôn ngữ thân thiện, dễ hiểu và cùng ngôn ngữ với ngôn ngữ của câu hỏi người dùng đưa ra

Ví dụ format câu trả lời:
- Điểm 1
- Điểm 2
- Điểm 3

Nguồn: [link hoặc tên website]
"""

class ChatSession:
    def __init__(self, session_id: str):
        self.session_id = session_id
        self.history = []
        self.session_file = os.path.join(SESSION_DIR, f"{session_id}.json")
        self.load_session()

    def load_session(self):
        if os.path.exists(self.session_file):
            with open(self.session_file, 'r', encoding='utf-8') as f:
                self.history = json.load(f)

    def save_session(self):
        os.makedirs(SESSION_DIR, exist_ok=True)
        with open(self.session_file, 'w', encoding='utf-8') as f:
            json.dump(self.history, f, ensure_ascii=False, indent=2)

    def add_interaction(self, question: str, answer: str):
        self.history.append({
            "timestamp": datetime.now().isoformat(),
            "question": question,
            "answer": answer
        })
        if len(self.history) > 5:  # Giữ lại 5 tương tác gần nhất
            self.history = self.history[-5:]
        self.save_session()

    def get_history_str(self) -> str:
        return "\n".join([
            f"Q: {interaction['question']}\nA: {interaction['answer']}"
            for interaction in self.history
        ])

def load_local_documents():
    documents = []
    logging.info(f"Đang tìm files trong thư mục: {os.path.abspath(LOCAL_DATA_DIR)}")
    
    if not os.path.exists(LOCAL_DATA_DIR):
        logging.error(f"Thư mục {LOCAL_DATA_DIR} không tồn tại")
        return documents
    
    def read_docx(file_path):
        doc = DocxDocument(file_path)
        return ' '.join([paragraph.text for paragraph in doc.paragraphs])
        
    def read_pdf(file_path):
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ''
            for page in pdf_reader.pages:
                text += page.extract_text() + '\n'
            return text
            
    def read_excel(file_path):
        df = pd.read_excel(file_path)
        return df.to_string()
    
    for root, dirs, files in os.walk(LOCAL_DATA_DIR):
        logging.info(f"Đang duyệt thư mục: {root}")
        logging.info(f"Các thư mục con: {dirs}")
        logging.info(f"Các files tìm thấy: {files}")
        
        for file in files:
            file_path = os.path.join(root, file)
            try:
                content = None
                
                if file.endswith('.txt'):
                    with open(file_path, 'rb') as f:
                        content = f.read().decode('utf-8')
                        
                elif file.endswith(('.doc', '.docx')):
                    content = read_docx(file_path)
                    
                elif file.endswith('.pdf'):
                    content = read_pdf(file_path)
                    
                elif file.endswith(('.xls', '.xlsx')):
                    content = read_excel(file_path)
                
                if content:
                    logging.info(f"Đã đọc file: {file_path}")
                    documents.append(
                        Document(
                            page_content=content,
                            metadata={"source": file_path}
                        )
                    )
                    
            except Exception as e:
                logging.error(f"Lỗi khi đọc file {file_path}: {e}")
    
    logging.info(f"Tổng số documents đã đọc: {len(documents)}")
    return documents

def split_text(documents):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    return text_splitter.split_documents(documents)

def create_vector_store(texts):
    if not texts:
        logging.error("Không có texts để upload")
        return None
        
    embeddings = OpenAIEmbeddings(
        model="text-embedding-ada-002",
        openai_api_key=OPENAI_API_KEY
    )
    
    logging.info(f"Bắt đầu tạo vector store với {len(texts)} documents")
    
    try:
        vector_store = Chroma.from_documents(
            documents=texts,
            embedding=embeddings,
            persist_directory=CHROMA_PERSIST_DIR
        )
        vector_store.persist()  # Lưu vector store
        logging.info("Tạo vector store thành công")
        return vector_store
    except Exception as e:
        logging.error(f"Lỗi khi tạo vector store: {e}")
        raise

def extract_keywords(question: str, top_n: int = 3) -> List[str]:
    kw_model = KeyBERT()
    keywords = kw_model.extract_keywords(question, keyphrase_ngram_range=(1, 2), stop_words='english', top_n=top_n)
    return [kw[0] for kw in keywords]

def search_on_web(query: str) -> Optional[str]:
    try:
        search = SerpAPIWrapper(serpapi_api_key=SERP_API_KEY)
        results = search.run(query)
        return results
    except Exception as e:
        logging.error(f"Lỗi khi tìm kiếm web: {e}")
        return None

def format_web_response(web_result: str, query: str, history: str = "") -> str:
    try:
        llm = OpenAI(model="gpt-3.5-turbo", openai_api_key=OPENAI_API_KEY)
        prompt = f"""
        {SYSTEM_PROMPT}
        
        Lịch sử hội thoại:
        {history}
        
        Câu hỏi hiện tại: "{query}"
        Thông tin tìm được: "{web_result}"
        
        Hãy format lại câu trả lời theo hướng dẫn trên.
        """
        
        formatted_response = llm.invoke(prompt)
        return formatted_response
    except Exception as e:
        logging.error(f"Lỗi khi format câu trả lời: {e}")
        return web_result

def initialize_chatbot(vector_store):
    llm = OpenAI(model="gpt-3.5-turbo", openai_api_key=OPENAI_API_KEY)
    retriever = vector_store.as_retriever(search_kwargs={"k": 5})
    
    chatbot = RetrievalQA.from_chain_type(
        llm=llm,
        chain_type="stuff",
        retriever=retriever,
        verbose=True
    )
    
    return chatbot

def chat(session_id: str = None):
    if not session_id:
        session_id = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    session = ChatSession(session_id)
    
    # Tải và xử lý dữ liệu local
    documents = load_local_documents()
    if not documents:
        print("Không có dữ liệu nào được tải. Vui lòng kiểm tra thư mục data/documents")
        return

    texts = split_text(documents)
    vector_store = create_vector_store(texts)
    chatbot = initialize_chatbot(vector_store)

    print("Chatbot: Xin chào! Bạn có thể đặt câu hỏi về tuyển sinh.")
    
    while True:
        try:
            user_input = input("Bạn: ").strip()
            if not user_input:
                print("Chatbot: Vui lòng nhập câu hỏi của bạn.")
                continue
                
            if user_input.lower() in ["e", "q", "tạm biệt", "exit", "quit"]:
                print("Chatbot: Tạm biệt! Hẹn gặp lại.")
                break

            # Lấy lịch sử hội thoại
            history = session.get_history_str()
            
            # Trích xuất từ khóa từ câu hỏi
            keywords = extract_keywords(user_input)
            if not keywords:
                print("Chatbot: Xin lỗi, tôi không hiểu câu hỏi của bạn. Vui lòng thử lại.")
                continue
            
            # Tìm kiếm trong vector store
            response = chatbot.invoke({"query": user_input})
            
            # Kiểm tra chất lượng câu trả lời và tìm kiếm web nếu cần
            if (not response or 
                len(str(response).strip()) < 50 or 
                "không tìm thấy thông tin" in str(response).lower() or
                "tôi không biết" in str(response).lower()):
                
                logging.info("Không tìm thấy thông tin trong context, đang tìm kiếm web...")
                print("Chatbot: Đang tìm kiếm thông tin bổ sung trên web...")
                
                search_query = f"tuyển sinh đại học {' '.join(keywords)}"
                web_result = search_on_web(search_query)
                
                if web_result:
                    response = format_web_response(web_result, user_input, history)
                else:
                    response = "Xin lỗi, tôi không tìm thấy thông tin liên quan cả trong cơ sở dữ liệu lẫn trên web."

            response_text = str(response)
            print(f"Chatbot: {response_text}")
            
            # Lưu tương tác vào session
            session.add_interaction(user_input, response_text)
            
        except KeyboardInterrupt:
            print("\nChatbot: Tạm biệt! Hẹn gặp lại.")
            break
        except Exception as e:
            logging.error(f"Lỗi không mong muốn: {e}")
            print("Chatbot: Xin lỗi, đã có lỗi xảy ra. Vui lòng thử lại.")

if __name__ == "__main__":
    chat() 