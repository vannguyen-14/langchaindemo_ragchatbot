from typing import List
from typing import Optional
import logging
import json
import os
from datetime import datetime
from langchain_community.embeddings import OpenAIEmbeddings
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain_community.chat_models import ChatOpenAI
from langchain_openai import OpenAIEmbeddings
from langchain.chains import RetrievalQA
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.utilities import SerpAPIWrapper
from keybert import KeyBERT
from docx import Document as DocxDocument
import PyPDF2
import pandas as pd
from langchain_core.documents import Document
from dotenv import load_dotenv  # Thêm import để sử dụng dotenv

# Tải các biến môi trường từ file .env
load_dotenv()

# Định nghĩa đường dẫn
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
LOCAL_DATA_DIR = os.path.join(BASE_DIR, "data", "documents")

# Cấu hình
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")  # Thay đổi để lấy từ biến môi trường
SERP_API_KEY = os.getenv("SERP_API_KEY")  # Thay đổi để lấy từ biến môi trường
CHROMA_PERSIST_DIR = "data/chroma_db"  # Thư mục lưu vector store

SYSTEM_PROMPT = """Bạn là một trợ lý tư vấn các thông tin về tuyển sinh, học phí, các thông tin về các trường đại học, v.v... thông minh và thân thiện. Khi trả lời:
1. Hãy tóm tắt thông tin một cách ngắn gọn, dễ hiểu
2. Trình bày thông tin theo dạng có cấu trúc với các bullet points
3. Nếu không tìm thấy thông tin trong các tài liệu được tải lên, hay tìm kiếm thông tin từ web, hãy trích dẫn nguồn và đưa vào dạng link
4. Tập trung vào những thông tin quan trọng và liên quan nhất
5. Sử dụng ngôn ngữ thân thiện, dễ hiểu và cùng ngôn ngữ với ngôn ngữ của câu hỏi người dùng đưa ra
6. Với những câu hỏi không liên quan tới lĩnh vực giáo dục nói chung và tuyển sinh nói riêng, hãy cố gắng hướng cuộc trò chuyện về lại lĩnh vực của bạn.

Ví dụ format câu trả lời:
- Điểm 1
- Điểm 2
- Điểm 3

Nguồn: [link hoặc tên website]
"""

# Hàm này dùng để tải tài liệu từ các file
def load_local_documents():
    documents = []
    logging.info(f"Đang tìm files trong thư mục: {os.path.abspath(LOCAL_DATA_DIR)}")

    if not os.path.exists(LOCAL_DATA_DIR):
        logging.error(f"Thư mục {LOCAL_DATA_DIR} không tồn tại")
        return documents

    def read_docx(file_path):
        doc = DocxDocument(file_path)
        return ' '.join([paragraph.text for paragraph in doc.paragraphs])

    def read_pdf(file_path):
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ''
            for page in pdf_reader.pages:
                text += page.extract_text() + '\n'
            return text

    def read_excel(file_path):
        df = pd.read_excel(file_path)
        return df.to_string()

    for root, dirs, files in os.walk(LOCAL_DATA_DIR):
        logging.info(f"Đang duyệt thư mục: {root}")
        logging.info(f"Các thư mục con: {dirs}")
        logging.info(f"Các files tìm thấy: {files}")

        for file in files:
            file_path = os.path.join(root, file)
            try:
                content = None

                if file.endswith('.txt'):
                    with open(file_path, 'rb') as f:
                        content = f.read().decode('utf-8')

                elif file.endswith(('.doc', '.docx')):
                    content = read_docx(file_path)

                elif file.endswith('.pdf'):
                    content = read_pdf(file_path)

                if content:
                    logging.info(f"Đã đọc file: {file_path}")
                    documents.append(
                        Document(
                            page_content=content,
                            metadata={"source": file_path}
                        )
                    )

            except Exception as e:
                logging.error(f"Lỗi khi đọc file {file_path}: {e}")

    logging.info(f"Tổng số documents đã đọc: {len(documents)}")
    return documents

# Chia nhỏ văn bản thành các đoạn nhỏ
def split_text(documents):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    return text_splitter.split_documents(documents)

# Tạo vector store từ các văn bản
def create_vector_store(texts):
    if not texts:
        logging.error("Không có texts để upload")
        return None

    embeddings = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)

    logging.info(f"Bắt đầu tạo vector store với {len(texts)} documents")

    try:
        vector_store = Chroma.from_documents(
            documents=texts,
            embedding=embeddings,
            persist_directory=CHROMA_PERSIST_DIR
        )
        vector_store.persist()  # Lưu vector store
        logging.info("Tạo vector store thành công")
        return vector_store
    except Exception as e:
        logging.error(f"Lỗi khi tạo vector store: {e}")
        raise

# Trích xuất từ khóa từ câu hỏi
def extract_keywords(question: str, top_n: int = 5) -> List[str]:
    kw_model = KeyBERT()
    keywords = kw_model.extract_keywords(question, keyphrase_ngram_range=(1, 2), stop_words='english', top_n=top_n)
    return [kw[0] for kw in keywords]

# Tìm kiếm trên web
def search_on_web(query: str) -> Optional[str]:
    try:
        search = SerpAPIWrapper(serpapi_api_key=SERP_API_KEY)
        results = search.run(query)
        return results
    except Exception as e:
        logging.error(f"Lỗi khi tìm kiếm web: {e}")
        return None

# Định dạng câu trả lời từ web
def format_web_response(web_result: str, query: str, history: str = "") -> str:
    try:
        llm = ChatOpenAI(model="gpt-4o-mini", openai_api_key=OPENAI_API_KEY)
        prompt = f"""
        {SYSTEM_PROMPT}
        
        Lịch sử hội thoại:
        {history}
        
        Câu hỏi hiện tại: "{query}"
        Thông tin tìm được: "{web_result}"
        
        Hãy format lại câu trả lời theo hướng dẫn trên.
        """

        formatted_response = llm.invoke(prompt)
        return formatted_response
    except Exception as e:
        logging.error(f"Lỗi khi format câu trả lời: {e}")
        return web_result

# Khởi tạo chatbot với vector store
def initialize_chatbot(vector_store):
    # Khởi tạo LLM OpenAI
    llm = ChatOpenAI(model="text-embedding-3-small", openai_api_key=OPENAI_API_KEY)

    # Tạo RetrievalQA chain sử dụng vector_store làm retriever
    qa_chain = RetrievalQA.from_chain_type(
        llm=llm,
        chain_type="stuff",  # Có thể thử các chain_type khác nếu cần
        retriever=vector_store.as_retriever()
    )

    logging.info("Chatbot đã được khởi tạo thành công với OpenAI.")
    return qa_chain

# Hàm trả lời câu hỏi từ chatbot
def get_answer(question, session_id=None):
    # Thiết lập và xử lý dữ liệu như load_local_documents(), split_text(), create_vector_store()
    documents = load_local_documents()  # Tải dữ liệu
    texts = split_text(documents)       # Chia nhỏ văn bản
    vector_store = create_vector_store(texts)  # Tạo vector store

    if not vector_store:
        return "Không thể tạo vector store, vui lòng kiểm tra lại dữ liệu."

    # Khởi tạo chatbot với vector store
    chatbot = initialize_chatbot(vector_store)

    # Tìm câu trả lời cho câu hỏi từ người dùng
    response = chatbot.invoke({"query": question})

    # Kiểm tra chất lượng câu trả lời và trả về
    if (not response or len(str(response).strip()) < 50 or "không tìm thấy thông tin" in str(response).lower()):
        # Nếu không có câu trả lời hợp lệ, tìm kiếm trên web
        search_query = f"tuyển sinh đại học {question}"
        web_result = search_on_web(search_query)

        if web_result:
            response = format_web_response(web_result, question)
        else:
            response = "Xin lỗi, tôi không tìm thấy thông tin liên quan cả trong cơ sở dữ liệu lẫn trên web."

    # Giữ lại chỉ phần nội dung câu trả lời
    response = response.get('result', '') if isinstance(response, dict) else response
    return response
